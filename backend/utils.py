import os
import io
import fitz  # PyMuPDF
from docx import Document
from werkzeug.utils import secure_filename

from langchain_community.document_loaders.blob_loaders import Blob
from langchain_core.documents import Document

def extract_lcdocs_from_file(file):
    filename = secure_filename(file.filename)
    file_extension = os.path.splitext(filename)[1]

    docs, text = None, ''
    if file_extension == '.txt':
        print('TEXT FILE')
        text = file.stream.read().decode('utf-8')
    elif file_extension == '.docx':
        print('DOCX FILE')
        doc = Document(io.BytesIO(file.stream.read()))
        text = ' '.join([paragraph.text for paragraph in doc.paragraphs])
    elif file_extension == '.pdf':
        print('PDF FILE')
        # getting lc docs to build rag chain
        docs = langchain_pdf_document_loader(file, filename)
    
    return docs, text



def langchain_pdf_document_loader(file, filename):
    blob = Blob.from_data(file.stream.read(), path=filename)
    with blob.as_bytes_io() as data:
        doc = fitz.open(stream=data, filetype="pdf")
        yield from [
            Document(page_content=page.get_text(),
                metadata=dict(
                    {
                        "source": blob.source,
                        "file_path": blob.source,
                        "page": page.number,
                        "total_pages": len(doc),
                    },
                    **{
                        k: doc.metadata[k]
                        for k in doc.metadata
                        if type(doc.metadata[k]) in [str, int]
                    },
                ),
            )
            for page in doc
        ]